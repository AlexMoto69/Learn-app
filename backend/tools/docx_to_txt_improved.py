r"""
Improved DOCX -> TXT converter.
Tries these methods in order:
  1. unzip XML and extract all <w:t> text nodes (best coverage for textboxes and shapes)
  2. docx2txt (handles text inside textboxes and images)
  3. mammoth (good for preserving structure and headings)
  4. python-docx fallback (extracts paragraphs, tables, headers/footers)

Usage (Windows cmd):
  pip install docx2txt mammoth python-docx
  python tools\docx_to_txt_improved.py "C:\path\to\biolaureat.docx" "lessons\biolaureat.txt"

This script writes UTF-8 text and returns exit code 0 on success.
"""
import sys
import os
from typing import Optional

import zipfile
import xml.etree.ElementTree as ET
import re


def write_text(dst_path: str, text: str) -> None:
    os.makedirs(os.path.dirname(dst_path) or ".", exist_ok=True)
    with open(dst_path, "w", encoding="utf-8") as f:
        f.write(text)


def try_unzip_xml(src: str) -> Optional[str]:
    """Unzip the .docx and extract text from all word/*.xml files by collecting <w:t> nodes.
    This often captures textboxes, headers/footers, footnotes, and other content python-docx misses.
    """
    try:
        with zipfile.ZipFile(src, 'r') as z:
            text_parts = []
            for name in z.namelist():
                if not name.startswith('word/') or not name.endswith('.xml'):
                    continue
                try:
                    data = z.read(name).decode('utf-8', errors='ignore')
                except Exception:
                    continue
                # Try XML parsing first; preserve paragraph boundaries by iterating <w:p>
                try:
                    root = ET.fromstring(data)
                    ns = ''
                    # detect namespace by root tag if present
                    if root.tag.startswith('{'):
                        ns = root.tag.split('}')[0] + '}'

                    paras = []
                    # find all paragraph elements
                    for p in root.findall('.//' + ns + 'p'):
                        texts = []
                        # collect all <w:t> under this paragraph
                        for t in p.findall('.//' + ns + 't'):
                            if t.text:
                                texts.append(t.text)
                        # handle line breaks inside paragraph
                        if texts:
                            para_text = ''.join(texts).strip()
                            if para_text:
                                paras.append(para_text)
                    if paras:
                        # join paragraphs with double newline to preserve distinct blocks
                        text = '\n\n'.join(paras)
                        text_parts.append(text)
                        continue
                except Exception:
                    # fallback to regex extraction
                    parts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', data, flags=re.DOTALL)
                    if parts:
                        text = '\n'.join(p.strip() for p in parts if p.strip())
                        text_parts.append(text)
            if text_parts:
                # Join parts from different xml files with double newline
                joined = '\n\n'.join(p.strip() for p in text_parts if p.strip())
                return joined
    except Exception:
        return None
    return None


def try_docx2txt(src: str) -> Optional[str]:
    try:
        import docx2txt
    except Exception:
        return None
    try:
        txt = docx2txt.process(src)
        if isinstance(txt, str) and txt.strip():
            return txt
        return None
    except Exception:
        return None


def try_mammoth(src: str) -> Optional[str]:
    try:
        import mammoth
    except Exception:
        return None
    try:
        with open(src, "rb") as docx_file:
            result = mammoth.convert_to_markdown(docx_file)
            md = result.value or ""
            if md.strip():
                return md
            return None
    except Exception:
        return None


def try_python_docx(src: str) -> Optional[str]:
    try:
        from docx import Document
    except Exception:
        return None
    try:
        doc = Document(src)
        parts = []

        # paragraphs
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

        # headers/footers from each section
        for section in doc.sections:
            header = section.header
            if header is not None:
                for p in header.paragraphs:
                    t = p.text.strip()
                    if t:
                        parts.append("[HEADER] " + t)
            footer = section.footer
            if footer is not None:
                for p in footer.paragraphs:
                    t = p.text.strip()
                    if t:
                        parts.append("[FOOTER] " + t)

        # tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join([c for c in cells if c])
                if row_text:
                    parts.append(row_text)

        out = "\n\n".join(parts).strip()
        return out if out else None
    except Exception:
        return None


def convert(src: str, dst: str) -> int:
    if not os.path.isfile(src):
        print(f"Source file not found: {src}")
        return 2

    # Try unzip/xml extraction first
    txt = try_unzip_xml(src)
    if txt:
        write_text(dst, txt)
        print(f"Converted using unzip-xml extractor -> {dst}")
        return 0

    # Try docx2txt
    txt = try_docx2txt(src)
    if txt:
        write_text(dst, txt)
        print(f"Converted using docx2txt -> {dst}")
        return 0

    # Try mammoth
    txt = try_mammoth(src)
    if txt:
        write_text(dst, txt)
        print(f"Converted using mammoth (markdown preserved) -> {dst}")
        return 0

    # Try python-docx fallback
    txt = try_python_docx(src)
    if txt:
        write_text(dst, txt)
        print(f"Converted using python-docx fallback -> {dst}")
        return 0

    print("Failed to convert document: no available converter or document empty")
    return 3


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python tools\\docx_to_txt_improved.py <input.docx> <output.txt>")
        sys.exit(1)
    src = sys.argv[1]
    dst = sys.argv[2]
    rc = convert(src, dst)
    sys.exit(rc)
